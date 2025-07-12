import io
from typing import Dict, Any, List
import PyPDF2
import pdfplumber
from docx import Document
import re

from .base_agent import BaseAgent, AgentResult

class DataExtractionAgent(BaseAgent):
    """
    Agent 1: Data Extraction Agent
    Responsible for ingesting raw resume files and extracting structured data
    """
    
    def __init__(self):
        super().__init__()
        self.agent_type = "data_extraction"
        self.capabilities = [
            "pdf_text_extraction",
            "docx_text_extraction", 
            "doc_text_extraction",
            "contact_info_extraction",
            "skills_identification",
            "experience_parsing"
        ]
        
    async def process(self, input_data: Dict[str, Any]) -> AgentResult:
        """
        Extract data from uploaded resume files
        
        Input: {
            "files": [{"filename": str, "content": bytes}],
            "job_id": str
        }
        
        Output: {
            "extracted_resumes": [
                {
                    "file_id": str,
                    "filename": str,
                    "raw_text": str,
                    "structured_data": {
                        "contact_info": {...},
                        "skills": [...],
                        "experience": [...],
                        "education": [...]
                    }
                }
            ]
        }
        """
        try:
            self.status = "processing"
            self.log_activity("Starting data extraction", {"file_count": len(input_data.get("files", []))})
            
            extracted_resumes = []
            
            for file_data in input_data.get("files", []):
                try:
                    # Extract text from file
                    raw_text = await self._extract_text_from_file(
                        file_data["content"], 
                        file_data["filename"]
                    )
                    
                    # Structure the extracted data
                    structured_data = await self._structure_extracted_data(raw_text)
                    
                    resume_data = {
                        "file_id": f"file_{len(extracted_resumes) + 1}",
                        "filename": file_data["filename"],
                        "raw_text": raw_text,
                        "structured_data": structured_data,
                        "extraction_metadata": {
                            "file_size": len(file_data["content"]),
                            "extraction_method": self._get_extraction_method(file_data["filename"]),
                            "text_length": len(raw_text),
                            "extraction_timestamp": self.created_at.isoformat()
                        }
                    }
                    
                    extracted_resumes.append(resume_data)
                    self.log_activity("File processed", {"filename": file_data["filename"]})
                    
                except Exception as e:
                    self.log_activity("File processing failed", {
                        "filename": file_data["filename"], 
                        "error": str(e)
                    })
                    continue
            
            self.status = "completed"
            
            return AgentResult(
                success=True,
                data={"extracted_resumes": extracted_resumes},
                metadata={
                    "total_files": len(input_data.get("files", [])),
                    "successful_extractions": len(extracted_resumes),
                    "agent_type": self.agent_type
                }
            )
            
        except Exception as e:
            self.status = "failed"
            self.log_activity("Data extraction failed", {"error": str(e)})
            return AgentResult(success=False, error=str(e))
    
    async def _extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """Extract text based on file type"""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return self._extract_from_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            return self._extract_from_docx(file_content)
        elif filename_lower.endswith('.doc'):
            return self._extract_from_doc(file_content)
        else:
            raise ValueError(f"Unsupported file format: {filename}")
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF using multiple methods"""
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                if text.strip():
                    return self._clean_text(text)
            
            # Fallback to PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def _extract_from_doc(self, file_content: bytes) -> str:
        """Extract text from DOC file (legacy format)"""
        try:
            # Try using python-docx (sometimes works with DOC)
            return self._extract_from_docx(file_content)
        except Exception as e:
            raise Exception(f"DOC extraction failed: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere
        text = re.sub(r'[^\w\s@.-]', ' ', text)
        
        # Normalize line breaks
        text = text.replace('\n', ' ').replace('\r', ' ')
        
        # Remove extra spaces
        text = ' '.join(text.split())
        
        return text.strip()
    
    async def _structure_extracted_data(self, raw_text: str) -> Dict[str, Any]:
        """Structure the extracted text into organized data"""
        
        # Extract contact information
        contact_info = self._extract_contact_info(raw_text)
        
        # Extract skills
        skills = self._extract_skills(raw_text)
        
        # Extract experience indicators
        experience = self._extract_experience_indicators(raw_text)
        
        # Extract education
        education = self._extract_education(raw_text)
        
        return {
            "contact_info": contact_info,
            "skills": skills,
            "experience": experience,
            "education": education
        }
    
    def _extract_contact_info(self, text: str) -> Dict[str, Any]:
        """Extract contact information using regex patterns"""
        contact_info = {}
        
        # Email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        contact_info["emails"] = emails[:3]  # Limit to first 3 emails
        
        # Phone extraction
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        contact_info["phones"] = [phone[0] + phone[1] if isinstance(phone, tuple) else phone for phone in phones[:2]]
        
        # Name extraction (simple heuristic)
        lines = text.split('\n')[:5]  # Check first 5 lines
        potential_names = []
        for line in lines:
            line = line.strip()
            if len(line.split()) == 2 and line.replace(' ', '').isalpha():
                potential_names.append(line)
        
        contact_info["potential_names"] = potential_names[:3]
        
        return contact_info
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills from text"""
        # Common technical skills database
        skill_keywords = [
            'python', 'java', 'javascript', 'react', 'angular', 'vue', 'node.js', 'express',
            'django', 'flask', 'spring', 'hibernate', 'sql', 'mysql', 'postgresql', 'mongodb',
            'redis', 'elasticsearch', 'docker', 'kubernetes', 'aws', 'azure', 'gcp',
            'git', 'jenkins', 'ci/cd', 'agile', 'scrum', 'html', 'css', 'sass', 'less',
            'typescript', 'php', 'ruby', 'go', 'rust', 'c++', 'c#', '.net', 'swift',
            'kotlin', 'flutter', 'react native', 'ios', 'android', 'machine learning',
            'artificial intelligence', 'data science', 'pandas', 'numpy', 'tensorflow',
            'pytorch', 'scikit-learn', 'tableau', 'power bi', 'excel', 'r', 'stata'
        ]
        
        text_lower = text.lower()
        found_skills = []
        
        for skill in skill_keywords:
            if skill in text_lower:
                found_skills.append(skill.title())
        
        return list(set(found_skills))  # Remove duplicates
    
    def _extract_experience_indicators(self, text: str) -> Dict[str, Any]:
        """Extract experience-related information"""
        
        # Look for years of experience
        experience_patterns = [
            r'(\d+)\+?\s*years?\s*of\s*experience',
            r'(\d+)\+?\s*years?\s*experience',
            r'experience\s*of\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*yrs?\s*experience'
        ]
        
        years_mentioned = []
        for pattern in experience_patterns:
            matches = re.findall(pattern, text.lower())
            years_mentioned.extend([int(match) for match in matches])
        
        # Look for job titles and companies
        job_title_keywords = [
            'developer', 'engineer', 'manager', 'analyst', 'consultant', 'specialist',
            'coordinator', 'administrator', 'architect', 'lead', 'senior', 'junior',
            'intern', 'associate', 'director', 'vice president', 'ceo', 'cto', 'cfo'
        ]
        
        found_titles = []
        text_lower = text.lower()
        for title in job_title_keywords:
            if title in text_lower:
                found_titles.append(title.title())
        
        return {
            "years_mentioned": years_mentioned,
            "max_years": max(years_mentioned) if years_mentioned else 0,
            "job_titles_found": list(set(found_titles))
        }
    
    def _extract_education(self, text: str) -> List[str]:
        """Extract education information"""
        education_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'mba', 'bs', 'ms', 'ba', 'ma',
            'computer science', 'engineering', 'business', 'mathematics', 'physics',
            'chemistry', 'biology', 'economics', 'finance', 'marketing', 'psychology'
        ]
        
        text_lower = text.lower()
        found_education = []
        
        for edu in education_keywords:
            if edu in text_lower:
                found_education.append(edu.title())
        
        return list(set(found_education))
    
    def _get_extraction_method(self, filename: str) -> str:
        """Get the extraction method used for the file"""
        if filename.lower().endswith('.pdf'):
            return "pdfplumber_pypdf2"
        elif filename.lower().endswith('.docx'):
            return "python_docx"
        elif filename.lower().endswith('.doc'):
            return "python_docx_legacy"
        else:
            return "unknown"