import PyPDF2
import pdfplumber
from docx import Document
import io
from typing import Optional
import re

class FileService:
    """Service for processing uploaded files"""
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            # Try with pdfplumber first (better for complex layouts)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                if text.strip():
                    return FileService._clean_text(text)
            
            # Fallback to PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return FileService._clean_text(text)
            
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return "Error: Could not extract text from PDF file"
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return FileService._clean_text(text)
            
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return "Error: Could not extract text from DOCX file"
    
    @staticmethod
    def extract_text_from_doc(file_content: bytes) -> str:
        """Extract text from DOC file (legacy format)"""
        try:
            # For DOC files, we'll try to use python-docx
            # Note: python-docx primarily supports DOCX, but sometimes works with DOC
            return FileService.extract_text_from_docx(file_content)
        except Exception as e:
            print(f"Error extracting text from DOC: {e}")
            return "Error: Could not extract text from DOC file. Please convert to PDF or DOCX format."
    
    @staticmethod
    def extract_text_from_file(file_content: bytes, filename: str) -> str:
        """Extract text from file based on extension"""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return FileService.extract_text_from_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            return FileService.extract_text_from_docx(file_content)
        elif filename_lower.endswith('.doc'):
            return FileService.extract_text_from_doc(file_content)
        else:
            return "Error: Unsupported file format. Please upload PDF, DOC, or DOCX files."
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with processing
        text = re.sub(r'[^\w\s@.-]', ' ', text)
        
        # Normalize line breaks
        text = text.replace('\n', ' ').replace('\r', ' ')
        
        # Remove extra spaces
        text = ' '.join(text.split())
        
        return text.strip()
    
    @staticmethod
    def validate_file(filename: str, file_size: int, max_size_mb: int = 10) -> tuple[bool, str]:
        """Validate uploaded file"""
        
        # Check file extension
        allowed_extensions = ['.pdf', '.doc', '.docx']
        filename_lower = filename.lower()
        
        if not any(filename_lower.endswith(ext) for ext in allowed_extensions):
            return False, "Invalid file format. Only PDF, DOC, and DOCX files are allowed."
        
        # Check file size (convert MB to bytes)
        max_size_bytes = max_size_mb * 1024 * 1024
        if file_size > max_size_bytes:
            return False, f"File size too large. Maximum allowed size is {max_size_mb}MB."
        
        # Check filename length
        if len(filename) > 255:
            return False, "Filename too long. Maximum 255 characters allowed."
        
        return True, "File is valid"